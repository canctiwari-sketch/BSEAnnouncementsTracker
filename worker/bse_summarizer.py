import requests
import os
import json
from datetime import datetime, timedelta
import PyPDF2
import io
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from web_search_utils import fetch_web_intel

load_dotenv()

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

api_key = os.getenv("GEMINI_API_KEY")

# Configuration
# You can get this URL by inspecting the network tab on bseindia.com -> Corporate Announcements
# This is a common endpoint, but it might change.
BSE_API_URL = "https://api.bseindia.com/BseIndiaAPI/api/AnnSubCategoryGetData/w?strCat=-1&strPrevDate={}&strScrip=&strSearch=P&strToDate={}&strType=C"
# Mocking a real browser is crucial to avoid 403 Forbidden
HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
    "Referer": "https://www.bseindia.com/",
    "Origin": "https://www.bseindia.com",
    "Accept": "application/json, text/plain, */*"
}

DOWNLOAD_DIR = "downloads"

def setup_directories():
    if not os.path.exists(DOWNLOAD_DIR):
        os.makedirs(DOWNLOAD_DIR)


def download_pdf(pdf_url, file_name):
    """Downloads a PDF file from a given URL."""
    try:
        response = requests.get(pdf_url, headers=HEADERS, stream=True)
        response.raise_for_status()

        file_path = os.path.join(DOWNLOAD_DIR, file_name)
        with open(file_path, "wb") as f:
            for chunk in response.iter_content(chunk_size=8192):
                f.write(chunk)
        print(f"Downloaded: {file_path}")
        return file_path
    except Exception as e:
        print(f"Error downloading PDF {pdf_url}: {e}")
        return None

def download_bse_attachment(attachment_name, file_name):
    """Tries multiple base URLs to find the PDF."""
    base_urls = [
        "https://www.bseindia.com/xml-data/corpfiling/AttachLive/",
        "https://www.bseindia.com/xml-data/corpfiling/AttachHis/"
    ]

    for base in base_urls:
        url = f"{base}{attachment_name}"
        path = download_pdf(url, file_name)
        if path:
            return path

    print(f"Failed to find PDF for {attachment_name} in Live or History.")
    return None

def extract_text_from_pdf(pdf_path):
    """Extracts text from a PDF file."""
    try:
        with open(pdf_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
        return text
    except Exception as e:
        print(f"Error extracting text from {pdf_path}: {e}")
        return ""

# Global variable to cache the working model name
GEMINI_MODEL = "gemini-2.5-flash-lite"  # Use same model as main fetch worker

def call_gemini(prompt, max_tokens=8192):
    """Call Gemini API directly via requests (no SDK dependency)."""
    if not api_key:
        return None
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent?key={api_key}"
    payload = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {"maxOutputTokens": max_tokens, "temperature": 0.3},
    }
    try:
        r = requests.post(url, json=payload, timeout=300)
        r.raise_for_status()
        data = r.json()
        return data["candidates"][0]["content"]["parts"][0]["text"]
    except Exception as e:
        print(f"Gemini API error: {e}")
        return None

def summarize_text(text):
    """Summarizes a single document using Gemini API."""
    if not text:
        return "No text to summarize."
    if not api_key:
        return "[MISSING CONFIG] GEMINI_API_KEY not set."
    prompt = f"""You are an expert Equity Research Analyst. Analyze this corporate announcement document.
Provide: 1) Executive Summary (2-3 sentences), 2) Key financial figures, 3) Future guidance/targets, 4) Risk factors.

Document:
{text[:50000]}"""
    result = call_gemini(prompt, max_tokens=2048)
    return result or "Error calling Gemini API"

def clean_text_for_pdf(text):
    """Replaces unsupported characters for standard PDF fonts."""
    replacements = {
        "₹": "Rs. ",
        "'": "'",
        "\u201c": '"',
        "\u201d": '"',
        "–": "-",
        "—": "-"
    }
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
    # Remove other non-latin-1 characters to prevent crashes
    return text.encode('latin-1', 'replace').decode('latin-1')

# ... imports ...
import json

def fetch_historical_announcements(scrip_code):
    """Fetches announcements for the last 3 years (2023-Present)."""
    announcements = []

    # Dynamic Date Generation (Quarterly to stay within API limits)
    start_year = 2023
    current_year = datetime.now().year
    today_str = datetime.now().strftime("%Y%m%d")

    periods = []
    for year in range(start_year, current_year + 1):
        # Quarterly chunks (Jan-Mar, Apr-Jun, Jul-Sep, Oct-Dec)
        chunks = [
            (f"{year}0101", f"{year}0331"),
            (f"{year}0401", f"{year}0630"),
            (f"{year}0701", f"{year}0930"),
            (f"{year}1001", f"{year}1231"),
        ]
        for start, end in chunks:
            if start > today_str:
                break
            if end > today_str:
                end = today_str
            periods.append((start, end))

    print(f"DEBUG: Generating report for period: {periods[0][0]} to {periods[-1][1]}")

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
        "Referer": "https://www.bseindia.com/",
        "Origin": "https://www.bseindia.com"
    }

    print("Fetching historical data (2023-2025)...")
    for start, end in periods:
        try:
            url = f"https://api.bseindia.com/BseIndiaAPI/api/AnnSubCategoryGetData/w?strCat=-1&strPrevDate={start}&strScrip={scrip_code}&strSearch=P&strToDate={end}&strType=C"
            response = requests.get(url, headers=headers)
            if response.status_code == 200:
                data = response.json()
                items = data.get("Table", [])
                print(f"  - {start} to {end}: Found {len(items)} items")
                announcements.extend(items)
        except Exception as e:
            print(f"Error fetching {start}-{end}: {e}")

    return announcements


def fetch_nse_announcements(nse_symbol):
    """Fetches announcements from NSE for the last 3 years."""
    if not nse_symbol:
        return []

    print(f"Fetching NSE announcements for {nse_symbol}...")
    announcements = []

    session = requests.Session()
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
        "Accept": "*/*",
        "Accept-Language": "en-US,en;q=0.9",
        "Referer": "https://www.nseindia.com/",
    }
    session.headers.update(headers)

    # Get cookies from NSE homepage
    try:
        session.get("https://www.nseindia.com", timeout=10)
    except:
        pass

    # Fetch in yearly chunks
    today = datetime.now()
    start_year = today.year - 3

    for year_offset in range(4):
        year = start_year + year_offset
        from_date = f"01-01-{year}"
        to_year = year
        to_date = f"31-12-{year}" if year < today.year else today.strftime("%d-%m-%Y")

        # Try both 'equities' and 'sme' index
        for index_type in ["equities", "sme"]:
            try:
                url = f"https://www.nseindia.com/api/corporate-announcements?index={index_type}&symbol={nse_symbol}&from_date={from_date}&to_date={to_date}"
                resp = session.get(url, timeout=15)
                if resp.status_code == 200:
                    data = resp.json()
                    if data:
                        print(f"  - NSE ({index_type}) {from_date} to {to_date}: Found {len(data)} items")
                        for item in data:
                            # Convert NSE format to match BSE format for compatibility
                            announcements.append({
                                "NEWSSUB": f"{item.get('sm_name', '')} - {item.get('desc', '')}",
                                "HEADNAME": item.get("attchmntText", ""),
                                "ATTACHMENTNAME": "",  # Not used for NSE
                                "NEWS_DT": item.get("dt", ""),
                                "NSE_PDF_URL": item.get("attchmntFile", ""),
                                "SOURCE": "NSE"
                            })
                        break  # Found data in this index type, skip other
            except Exception as e:
                print(f"  NSE Error ({index_type}) {year}: {e}")

    print(f"  NSE Total: {len(announcements)} announcements")
    return announcements


def is_important_document(subject, description=""):
    """Filters for high-value documents by checking both subject and description."""
    combined_text = f"{subject} {description}".lower()
    keywords = [
        "investor presentation", "earnings call", "transcript", "financial results",
        "annual report", "outcome of board meeting", "capacity", "expansion",
        "new product", "launch", "acquisition", "merger", "capex", "order", "award",
        "investor meet", "analyst meet", "conference call", "concall", "audio visual",
        "sast", "insider", "promoter", "pledge", "regulation 29", "regulation 7"
    ]
    exclude = ["trading window", "certificate", "duplicate", "agm", "scrutinizer", "loss of share", "intimation"]

    if any(ex in combined_text for ex in exclude):
        return False
    return any(k in combined_text for k in keywords)

def fetch_insider_trading(scrip_code):
    """Fetches insider trading data from BSE for the last 3 years. Returns summary of Market Purchase and Market Sale."""
    print(f"  [Insider] Fetching insider trading data for {scrip_code}...")

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
        "Referer": "https://www.bseindia.com/",
        "Origin": "https://www.bseindia.com",
        "Accept": "application/json, text/plain, */*"
    }

    today = datetime.now()
    from_date = (today - timedelta(days=3*365)).strftime("%Y%m%d")
    to_date = today.strftime("%Y%m%d")

    try:
        url = f"https://api.bseindia.com/BseIndiaAPI/api/InsiderTrading/w?scripcode={scrip_code}&fromdate={from_date}&todate={to_date}"
        response = requests.get(url, headers=headers)
        if response.status_code != 200:
            print(f"  [Insider] API returned status {response.status_code}")
            return ""

        data = response.json()
        if not data:
            print("  [Insider] No insider trading data found.")
            return ""

        market_purchase_shares = 0
        market_purchase_value = 0.0
        market_sale_shares = 0
        market_sale_value = 0.0

        for record in data:
            mode = str(record.get("MODOFACQ", "")).strip().lower()
            txn_type = str(record.get("TRANSACTIONTYPE", "")).strip().lower()

            if "market" not in mode:
                continue

            try:
                num_shares = int(str(record.get("NOOFSHARE", "0")).replace(",", ""))
            except (ValueError, TypeError):
                num_shares = 0
            try:
                value = float(str(record.get("VALUE", "0")).replace(",", ""))
            except (ValueError, TypeError):
                value = 0.0

            if "purchase" in mode or "buy" in txn_type or "acquisition" in txn_type:
                market_purchase_shares += num_shares
                market_purchase_value += value
            elif "sale" in mode or "sell" in txn_type or "disposal" in txn_type:
                market_sale_shares += num_shares
                market_sale_value += value

        if market_purchase_shares == 0 and market_sale_shares == 0:
            print("  [Insider] No market purchase/sale transactions found.")
            return ""

        summary = "### INSIDER TRADING SUMMARY (Last 3 Years - Market Transactions Only) ###\n\n"
        summary += f"**Market Purchases:** {market_purchase_shares:,} shares | Total Value: Rs. {market_purchase_value:,.0f}\n"
        summary += f"**Market Sales:** {market_sale_shares:,} shares | Total Value: Rs. {market_sale_value:,.0f}\n"
        net_shares = market_purchase_shares - market_sale_shares
        net_value = market_purchase_value - market_sale_value
        summary += f"**Net Position:** {'+' if net_shares >= 0 else ''}{net_shares:,} shares | Net Value: Rs. {'+' if net_value >= 0 else ''}{net_value:,.0f}\n"

        print(f"  [Insider] Found: Purchases={market_purchase_shares:,} shares, Sales={market_sale_shares:,} shares")
        return summary

    except Exception as e:
        print(f"  [Insider] Error fetching insider trading: {e}")
        return ""


def fetch_screener_financials(stock_name):
    """
    Fetches consolidated P&L, Balance Sheet, and Cash Flow data from screener.in.
    Returns a dict with:
      - 'profit-loss': markdown table string
      - 'balance-sheet': markdown table string
      - 'cash-flow': markdown table string
      - 'raw_text': combined text for Gemini context (summary only)
    """
    print(f"  [Screener] Fetching financial data for {stock_name}...")

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
    }

    result = {"profit-loss": "", "balance-sheet": "", "cash-flow": "", "raw_text": "", "metrics": {}}

    try:
        # Step 1: Search for the company to get the URL slug
        search_url = f"https://www.screener.in/api/company/search/?q={requests.utils.quote(stock_name)}"
        resp = requests.get(search_url, headers=headers)
        if resp.status_code != 200:
            print(f"  [Screener] Search failed with status {resp.status_code}")
            return result

        search_results = resp.json()
        if not search_results:
            print(f"  [Screener] No company found for '{stock_name}'")
            return result

        company_url = search_results[0].get("url", "")
        if not company_url:
            return result

        # Prefer consolidated
        if "/consolidated/" not in company_url:
            company_url = company_url.rstrip("/") + "/consolidated/"

        full_url = f"https://www.screener.in{company_url}"
        print(f"  [Screener] Fetching from {full_url}")

        # Step 2: Fetch the company page
        resp = requests.get(full_url, headers=headers)
        if resp.status_code != 200:
            print(f"  [Screener] Page fetch failed with status {resp.status_code}")
            return result

        soup = BeautifulSoup(resp.text, 'html.parser')

        def extract_table_markdown(section_tag):
            """Extract a clean markdown table from a screener section."""
            table = section_tag.find('table')
            if not table:
                return ""

            lines = []
            thead = table.find('thead')
            if thead:
                header_cells = thead.find_all('th')
                headers_text = [cell.get_text(strip=True) for cell in header_cells]
                lines.append("| " + " | ".join(headers_text) + " |")
                lines.append("| " + " | ".join(["---"] * len(headers_text)) + " |")

            tbody = table.find('tbody')
            if tbody:
                for row in tbody.find_all('tr'):
                    cells = row.find_all(['td', 'th'])
                    row_data = [cell.get_text(strip=True) for cell in cells]
                    lines.append("| " + " | ".join(row_data) + " |")

            return "\n".join(lines)

        # Extract key metrics from the top of the page
        metrics = {}
        # Screener uses <li> items with <span class="name"> and <span class="number">
        # or #top-ratios section with list items
        top_ratios = soup.find('div', id='top-ratios') or soup.find('ul', id='top-ratios')
        if not top_ratios:
            # Try finding all li elements with ratio data
            top_ratios = soup.find('div', class_='company-ratios')

        # Parse all ratio items on the page
        ratio_items = soup.find_all('li', class_='flex')
        if not ratio_items:
            ratio_items = soup.select('#top-ratios li') or soup.select('.ratios-table li') or []

        # Broader approach: find all span pairs with name/number
        all_spans = soup.find_all('span', class_='name')
        for span in all_spans:
            name = span.get_text(strip=True).rstrip(':')
            # Find the next sibling or parent's number span
            number_span = span.find_next_sibling('span', class_='number') or span.find_next('span', class_='number')
            if number_span:
                value = number_span.get_text(strip=True)
                metrics[name.lower()] = value

        # Also try parsing from raw text patterns
        page_text = soup.get_text()
        import re as re_mod

        metric_patterns = {
            'market_cap': r'Market Cap[:\s]*₹?\s*([\d,]+(?:\.\d+)?)\s*Cr',
            'current_price': r'Current Price[:\s]*₹?\s*([\d,]+(?:\.\d+)?)',
            'stock_pe': r'Stock P/E[:\s]*([\d,]+(?:\.\d+)?)',
            'book_value': r'Book Value[:\s]*₹?\s*([\d,]+(?:\.\d+)?)',
            'face_value': r'Face Value[:\s]*₹?\s*([\d,]+(?:\.\d+)?)',
            'high_low': r'High / Low[:\s]*₹?\s*([\d,]+(?:\.\d+)?)\s*/\s*₹?\s*([\d,]+(?:\.\d+)?)',
            'roce': r'ROCE[:\s]*([\d,]+(?:\.\d+)?)\s*%',
            'roe': r'ROE[:\s]*([\d,]+(?:\.\d+)?)\s*%',
            'dividend_yield': r'Dividend Yield[:\s]*([\d,]+(?:\.\d+)?)\s*%',
        }

        for key, pattern in metric_patterns.items():
            match = re_mod.search(pattern, page_text)
            if match:
                if key == 'high_low':
                    metrics[key] = f"₹{match.group(1)} / ₹{match.group(2)}"
                else:
                    metrics[key] = match.group(1)

        # Calculate Price to Sales (Market Cap / TTM Sales)
        try:
            market_cap_str = metrics.get('market_cap', '').replace(',', '')
            if market_cap_str:
                market_cap = float(market_cap_str)
                # Get TTM sales from P&L section
                pl_section = soup.find('section', id='profit-loss')
                if pl_section:
                    pl_table = pl_section.find('table')
                    if pl_table:
                        # TTM is usually the last column
                        thead = pl_table.find('thead')
                        if thead:
                            headers_list = [th.get_text(strip=True) for th in thead.find_all('th')]
                            ttm_idx = None
                            for i, h in enumerate(headers_list):
                                if 'TTM' in h:
                                    ttm_idx = i
                                    break
                            if ttm_idx is None:
                                ttm_idx = len(headers_list) - 1  # Last column as fallback

                            # Get first data row (Sales)
                            tbody = pl_table.find('tbody')
                            if tbody:
                                first_row = tbody.find('tr')
                                if first_row:
                                    cells = first_row.find_all(['td', 'th'])
                                    if ttm_idx < len(cells):
                                        ttm_sales_str = cells[ttm_idx].get_text(strip=True).replace(',', '')
                                        try:
                                            ttm_sales = float(ttm_sales_str)
                                            if ttm_sales > 0:
                                                price_to_sales = round(market_cap / ttm_sales, 2)
                                                metrics['price_to_sales'] = str(price_to_sales)
                                        except ValueError:
                                            pass
        except Exception as e:
            print(f"  [Screener] P/S calculation error: {e}")

        # Calculate P/B from current price and book value
        try:
            price_str = metrics.get('current_price', '').replace(',', '')
            bv_str = metrics.get('book_value', '').replace(',', '')
            if price_str and bv_str:
                price = float(price_str)
                bv = float(bv_str)
                if bv > 0:
                    metrics['price_to_book'] = str(round(price / bv, 2))
        except Exception:
            pass

        result["metrics"] = metrics
        print(f"  [Screener] Metrics: CMP=₹{metrics.get('current_price', 'N/A')}, MCap=₹{metrics.get('market_cap', 'N/A')}Cr, P/E={metrics.get('stock_pe', 'N/A')}, P/B={metrics.get('price_to_book', 'N/A')}, P/S={metrics.get('price_to_sales', 'N/A')}")

        # Extract all data sections
        sections = soup.find_all('section')
        for section in sections:
            section_id = section.get('id', '')
            if section_id in ('profit-loss', 'balance-sheet', 'cash-flow'):
                table_md = extract_table_markdown(section)
                if table_md:
                    result[section_id] = table_md

        # Create context for Gemini with the metrics
        metrics_text = f"""### SCREENER.IN DATA ###
Key Metrics (from screener.in):
- CMP: ₹{metrics.get('current_price', 'N/A')}
- Market Cap: ₹{metrics.get('market_cap', 'N/A')} Cr
- Stock P/E (TTM): {metrics.get('stock_pe', 'N/A')}
- Price to Book: {metrics.get('price_to_book', 'N/A')}x
- Price to Sales: {metrics.get('price_to_sales', 'N/A')}x
- Book Value: ₹{metrics.get('book_value', 'N/A')}
- 52-Week High/Low: {metrics.get('high_low', 'N/A')}
- ROCE: {metrics.get('roce', 'N/A')}%
- ROE: {metrics.get('roe', 'N/A')}%
- Dividend Yield: {metrics.get('dividend_yield', 'N/A')}%

Financial tables (P&L, Balance Sheet, Cash Flow) will be inserted directly into the report. You do NOT need to generate these tables.
"""
        result["raw_text"] = metrics_text

        print(f"  [Screener] Successfully fetched financial data")
        return result

    except Exception as e:
        print(f"  [Screener] Error fetching financials: {e}")
        return result


def get_collective_summary(stock_name, collected_texts):
    """
    Generates a DEEP DIVE 15-page style report.
    """
    if not collected_texts:
        return "No significant documents found."

    # Reverse to put the newest documents (2025/2026) at the TOP of the prompt
    collected_texts.reverse()
    full_text = "\n\n".join(collected_texts)

    try:
        current_date_str = datetime.now().strftime("%B %Y")

        prompt = f"""
You are a Senior Equity Research Analyst at a top-tier institutional firm writing a **Comprehensive Company Analysis** report.

**Current Context:** Today is {current_date_str}.
- **FY25 (Apr'24-Mar'25) is OVER.** Treat FY25 numbers as **ACTUALS** (History).
- **FY26 (Apr'25-Mar'26) is the CURRENT Year.** (Q4 is ongoing).
- **FY27** and beyond are the Future.

Create a **Comprehensive Company Analysis** for **{stock_name}**.

Source Material:
{full_text[:800000]}

**CRITICAL INSTRUCTIONS:**
1. **Prioritize Recent Data:** Focus heavily on **FY25 and FY26** data. Read thoroughly the Earnings Call Transcripts, Investor Presentations, and Annual Reports.
2. **No Citations:** Do NOT include any source citations, references, or footnotes like *[Q3 FY26 Investor Presentation, Slide 9]* or *[AR FY25, Page 172]*. Write cleanly without any bracketed references.
3. **Logic:** For every estimate, provide a **short 1-line logic**.
4. **Tables:** Generate Markdown tables where specified. Tables must be clean with | separators.
5. **Depth:** This must be an exhaustive 10-15 page report. Do NOT be brief. Each section must be detailed and substantive.
6. **Numbers:** Extract and present ALL available financial data. Do not summarize away the numbers.
7. **Financial Tables:** Do NOT generate P&L, Balance Sheet, or Cash Flow tables in Section 3. These will be auto-inserted from screener.in. Just write a brief financial commentary.

Structure your report EXACTLY as follows (ALL 10 sections are MANDATORY):

# {stock_name}
## Comprehensive Company Analysis
### Investment Research Note

## 1. Executive Summary
Write a detailed 2-3 paragraph executive summary covering:
- What the company does and its current position
- Key transformative events in the past 1-2 years
- Post-transformation metrics (revenue base, debt status, growth targets)
- Key growth catalysts ahead

Then include a **Snapshot** table. The screener.in data section has exact values for CMP, Market Cap, P/E, P/B, P/S, 52-Week range, ROE, ROCE — use those EXACT values, do NOT leave them as "...":
| Metric | Value |
|---|---|
| CMP | (use current_price from screener data) |
| Market Cap | (use market_cap from screener data, in Cr) |
| P/E (TTM) | (use stock_pe from screener data) |
| P/B | (use price_to_book from screener data) |
| P/S (TTM) | (use price_to_sales from screener data) |
| 52-Week High/Low | (use high_low from screener data) |
| Debt-Equity Ratio | ... |
| ROE | (use roe from screener data) |
| ROCE | (use roce from screener data) |
| Promoter Holding | ... |
| Dividend Yield | (use dividend_yield from screener data) |

## 1A. Company Brief
Write a detailed paragraph covering:
- What the company manufactures / what services it provides
- Key raw materials used and finished goods produced
- Domestic vs Export revenue mix (percentage split if available)
- Installed capacities across plants/segments and current utilisation levels
- Key end-user industries served
- Number and location of manufacturing facilities
This should read as a comprehensive introduction for someone unfamiliar with the company.

## 1B. News & Industry Overview
Write 2-3 detailed paragraphs covering:
- Recent news about the company from the last 6-12 months (management interviews, media coverage, analyst commentary)
- Industry trends, tailwinds, or headwinds affecting the company's sector
- Competitive landscape and where this company stands relative to peers
- Any regulatory changes or government policies impacting the business
Present this as flowing paragraphs (NOT bullet points).

## 2. Key Developments Over Past 1 Year
Cover ALL major developments with sub-sections (use ### for each). Include:
- Major divestitures, acquisitions, mergers
- New listings or exchange approvals
- Balance sheet transformation (include before/after table if applicable)
- Leadership changes
- Preferential issues, warrants, capital raises
- Any SAST disclosures, promoter buying/selling
- Any regulatory actions or penalties

For each development, provide dates and specifics.

## 3. Financial Analysis
**IMPORTANT: Do NOT generate P&L, Balance Sheet, or Cash Flow tables.** These will be inserted automatically from screener.in data. Just write the section heading "## 3. Financial Analysis" and then write a brief 2-3 line commentary on the overall financial health and trends, then move on to Section 4.
| Closing Cash & Equivalents | ... | ... |

Add commentary on cash flow drivers.


## 4. Corporate Governance Assessment

### 4.1 Board Composition
| Name | Designation | Category |
|---|---|---|
| ... | ... | ... |

### 4.2 Governance Strengths
List all positives (auditor opinion, internal controls, compliance, whistle-blower, etc.)

### 4.3 Governance Concerns / Watchpoints
List all red flags (family dominance, high remuneration ratios, penalties, related party loans, no dividends, etc.)

### 4.4 Shareholding Pattern
| Category | No. of Shares | % Holding |
|---|---|---|
| Promoters & PAIC | ... | ... |
| Banks, FIs, Mutual Funds | ... | ... |
| FIIs, NRIs & FPI | ... | ... |
| Private Corporate Bodies | ... | ... |
| Indian Public & Others | ... | ... |
| Total | ... | 100.00% |

Add notes on effective promoter group holding and any notable shareholders.

### 4.5 Insider Trading Activity (Last 3 Years)
From the SAST disclosures and insider trading filings in the source material, extract and summarize ONLY **Market Purchase** and **Market Sale** transactions. Ignore ESOP, Gift, Off Market, and other modes. Present a summary table:
| Transaction Type | Total Shares | Total Value (Rs.) |
|---|---|---|
| Market Purchases | ... | ... |
| Market Sales | ... | ... |
| **Net Position** | **...** | **...** |

Add a note on whether insiders have been net buyers or net sellers, and what this signals about management confidence.

## 5. Company & Business Quality Assessment

### 5.1 Business Model
Describe the business model in detail. Cover each business vertical/segment with revenue contribution percentages, key markets, capacity details, and growth rates.

### 5.2 Competitive Advantages (Moats)
List ALL competitive advantages with specifics:
- Manufacturing scale/capacity
- Regulatory approvals and certifications
- Key partnerships and contracts
- Balance sheet strength
- R&D pipeline
- Any hidden value (stakes in other companies, real estate, etc.)

### 5.3 Risks & Weaknesses
List ALL risks comprehensively:
- Execution risk
- Customer/revenue concentration
- Margin pressure
- Working capital intensity
- Governance concerns
- Competition and scale vs peers
- Currency/commodity risks
- Legacy issues

## 6. Growth Outlook & Strategy

### 6.1 Management's Revenue Roadmap
| Growth Driver | Current (FY25/26) | Target (FY28-29E) | CAGR |
|---|---|---|---|
| ... | ... | ... | ... |

### 6.2-6.5 Detail each growth driver
For each major growth initiative, write a dedicated sub-section explaining the strategy, timeline, and expected impact.

## 7. Future Revenue & EBITDA Estimates
| Particulars | FY25A | FY26E | FY27E | FY28E | FY29E |
|---|---|---|---|---|---|
| Revenue from Operations | ... | ... | ... | ... | ... |
| Revenue Growth (%) | ... | ... | ... | ... | ... |
| Gross Profit | ... | ... | ... | ... | ... |
| Gross Margin (%) | ... | ... | ... | ... | ... |
| EBITDA (Operating) | ... | ... | ... | ... | ... |
| EBITDA Margin (%) | ... | ... | ... | ... | ... |
| Depreciation | ... | ... | ... | ... | ... |
| Finance Cost | ... | ... | ... | ... | ... |
| PBT (Operating) | ... | ... | ... | ... | ... |
| Tax | ... | ... | ... | ... | ... |
| PAT (Operating) | ... | ... | ... | ... | ... |
| EPS (Operating) | ... | ... | ... | ... | ... |

### 7.1 Key Assumptions Behind Estimates
Explain assumptions for Revenue, Gross Margin, EBITDA Margin, Capex, Working Capital, Tax Rate, and Share Count.

## 8. Capex & Management Guidance Summary
| Guidance Parameter | Target | Timeline |
|---|---|---|
| Revenue Target | ... | ... |
| Revenue CAGR | ... | ... |
| Key Segment Targets | ... | ... |
| EBITDA Margin Target | ... | ... |
| Capex Plans | ... | ... |
| Dividend Policy | ... | ... |
| Debt Status | ... | ... |

## 9. Conclusion & Investment Summary

### Bull Case
Write a substantive paragraph on why this stock could outperform.

### Bear Case
Write a substantive paragraph on key risks that could derail the thesis.

### Key Monitorables
List 6-10 specific items investors should track going forward.

**Disclaimer:** This analysis is for informational purposes only and does not constitute investment advice. Please consult a registered financial advisor before making investment decisions.
"""

        result = call_gemini(prompt, max_tokens=16000)
        return result or "Error creating deep dive: no response from Gemini"
    except Exception as e:
        return f"Error creating deep dive: {e}"

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

# Professional color scheme (matching Ind-Swift style)
COLOR_PRIMARY = RGBColor(0x1F, 0x5F, 0x8B)    # Steel Blue for headings
COLOR_HEADER_BG = "1F5F8B"                      # Table header background
COLOR_ALT_ROW = "E8F0FE"                        # Alternating row background
COLOR_TEXT = RGBColor(0x33, 0x33, 0x33)          # Dark gray body text
COLOR_SUBTLE = RGBColor(0x66, 0x66, 0x66)        # Lighter gray for subtitles


def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def format_cell_text(cell, text, bold=False, font_size=9, font_color=None, alignment=None):
    """Format text within a table cell."""
    for paragraph in cell.paragraphs:
        paragraph.clear()
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Calibri'
    if bold:
        run.bold = True
    if font_color:
        run.font.color.rgb = font_color


def add_markdown_paragraph(doc, text, style='Normal'):
    """Parses simple markdown (**bold**) and adds a paragraph."""
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)

    # Simple bold parser: splits by **
    parts = text.split('**')
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_TEXT
        if i % 2 == 1:  # Odd parts were inside ** **
            run.bold = True
    return p


def style_heading(heading, level):
    """Apply professional styling to headings."""
    for run in heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = COLOR_PRIMARY
        if level == 0:
            run.font.size = Pt(26)
        elif level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(13)


def is_separator_row(cells):
    """Check if a table row is a markdown separator (e.g., |---|---|)."""
    return all(re.match(r'^[\s\-:]+$', c.strip()) for c in cells if c.strip())


def add_professional_table(doc, table_lines):
    """Create a professionally styled table from markdown table lines."""
    try:
        rows = []
        for line in table_lines:
            if not line.strip():
                continue
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            # Skip separator rows (---|---|---)
            if is_separator_row(cells):
                continue
            rows.append(cells)

        if not rows or len(rows) < 2:
            return

        max_cols = max(len(r) for r in rows)

        # Pad rows with fewer columns
        for r in rows:
            while len(r) < max_cols:
                r.append('')

        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        # Style the header row
        for c_idx in range(max_cols):
            cell = table.cell(0, c_idx)
            cell_text = rows[0][c_idx] if c_idx < len(rows[0]) else ''
            set_cell_shading(cell, COLOR_HEADER_BG)
            format_cell_text(cell, cell_text, bold=True, font_size=9,
                           font_color=RGBColor(0xFF, 0xFF, 0xFF))

        # Style data rows with alternating colors
        for r_idx in range(1, len(rows)):
            for c_idx in range(max_cols):
                cell = table.cell(r_idx, c_idx)
                cell_text = rows[r_idx][c_idx] if c_idx < len(rows[r_idx]) else ''

                # Alternating row shading
                if r_idx % 2 == 0:
                    set_cell_shading(cell, COLOR_ALT_ROW)

                # Bold the first column (labels)
                is_first_col = (c_idx == 0)
                format_cell_text(cell, cell_text, bold=is_first_col, font_size=9)

        # Set table borders
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            '  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            '  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            '  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            '</w:tblBorders>'
        )
        tblPr.append(borders)

        # Add some spacing after table
        doc.add_paragraph('')

    except Exception as e:
        print(f"Error creating professional table: {e}")


def add_cover_page(doc, stock_name):
    """Add a professional cover page."""
    # Add spacing before title
    for _ in range(6):
        p = doc.add_paragraph('')
        p.paragraph_format.space_after = Pt(0)

    # Company name - large and bold
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(stock_name.upper())
    run.font.size = Pt(32)
    run.font.name = 'Calibri'
    run.font.color.rgb = COLOR_PRIMARY
    run.bold = True

    # Divider line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('_' * 60)
    run.font.color.rgb = COLOR_PRIMARY
    run.font.size = Pt(8)

    # Subtitle - Comprehensive Company Analysis
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run('Comprehensive Company Analysis')
    run.font.size = Pt(22)
    run.font.name = 'Calibri'
    run.font.color.rgb = COLOR_PRIMARY

    # Sub-subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run('Investment Research Note')
    run.font.size = Pt(14)
    run.font.name = 'Calibri'
    run.font.color.rgb = COLOR_SUBTLE

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(48)
    run = p.add_run(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    run.font.color.rgb = COLOR_SUBTLE


    # Page break after cover
    doc.add_page_break()


def generate_word_report(stock_name, report_text):
    """Generates a professional Word report with institutional-grade formatting."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = COLOR_TEXT

    # Style headings
    for i in range(4):
        heading_style = doc.styles[f'Heading {i+1}'] if i > 0 else doc.styles['Title']
        hfont = heading_style.font
        hfont.name = 'Calibri'
        hfont.color.rgb = COLOR_PRIMARY

    # Add professional cover page
    add_cover_page(doc, stock_name)

    # Add header/footer
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run(f'{stock_name} | Comprehensive Analysis')
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_SUBTLE
    run.font.name = 'Calibri'

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run('Disclaimer: For informational purposes only. Not investment advice.')
    run.font.size = Pt(7)
    run.font.color.rgb = COLOR_SUBTLE
    run.font.name = 'Calibri'

    # Parse report content
    lines = report_text.split('\n')
    table_lines = []
    skip_cover_headers = True  # Skip the first few lines that duplicate cover page info

    for line in lines:
        stripped = line.strip()

        # Skip lines that are part of the cover page info (already rendered)
        if skip_cover_headers:
            if stripped.startswith('# ') and not stripped.startswith('## '):
                continue  # Skip top-level title (already on cover)
            if 'Comprehensive Company Analysis' in stripped:
                continue
            if 'Investment Research Note' in stripped:
                continue
            if stripped.startswith('## 1.'):
                skip_cover_headers = False  # Start processing from section 1

        # Table Detection & Processing
        if "|" in stripped and len(stripped) > 5:
            table_lines.append(stripped)
            continue
        elif table_lines:
            add_professional_table(doc, table_lines)
            table_lines = []

        if not stripped:
            continue

        # Headers
        if stripped.startswith('### '):
            heading_text = stripped.lstrip('#').strip()
            heading = doc.add_heading(heading_text, 2)
            style_heading(heading, 2)
        elif stripped.startswith('## '):
            heading_text = stripped.lstrip('#').strip()
            heading = doc.add_heading(heading_text, 1)
            style_heading(heading, 1)
            # Add a subtle line under major sections
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
        elif stripped.startswith('# '):
            heading_text = stripped.lstrip('#').strip()
            heading = doc.add_heading(heading_text, 0)
            style_heading(heading, 0)

        # Numbered lists (1. 2. 3. etc.)
        elif re.match(r'^\d+\.\s+', stripped):
            text = re.sub(r'^\d+\.\s+', '', stripped)
            add_markdown_paragraph(doc, text, style='List Number')

        # Bullet lists
        elif stripped.startswith('- ') or stripped.startswith('* '):
            bullet_text = stripped[2:]
            # Check for bold prefix pattern like "**Label:** description"
            add_markdown_paragraph(doc, bullet_text, style='List Bullet')

        # Horizontal rules / Disclaimer
        elif stripped.startswith('---'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)

        # Italic text (notes, disclaimers)
        elif stripped.startswith('*') and stripped.endswith('*') and not stripped.startswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip('*'))
            run.font.name = 'Calibri'
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = COLOR_SUBTLE

        # Normal Text
        else:
            add_markdown_paragraph(doc, stripped)

    # Process any remaining table lines
    if table_lines:
        add_professional_table(doc, table_lines)

    # Set margins for the document
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    safe_filename = f"{stock_name.replace(' ', '_')}_Analysis_Report.docx"
    output_file = os.path.join(ROOT, "data", "research", safe_filename)

    try:
        doc.save(output_file)
        print(f"\n[SUCCESS] Word Report saved to: {output_file}")
    except PermissionError:
        print(f"\n[WARNING] Could not write to {output_file} (File might be open).")
        base_filename = f"{stock_name.replace(' ', '_')}_Analysis_Report"
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_filename = f"{base_filename}_{timestamp}.docx"
        output_file = os.path.join(ROOT, "data", "research", safe_filename)
        print(f"Saving to new file instead: {output_file}")
        doc.save(output_file)

    return output_file

def download_nse_pdf(pdf_url, file_name):
    """Downloads a PDF from NSE archives."""
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            "Referer": "https://www.nseindia.com/",
        }
        response = requests.get(pdf_url, headers=headers, stream=True, timeout=30)
        response.raise_for_status()

        file_path = os.path.join(DOWNLOAD_DIR, file_name)
        with open(file_path, "wb") as f:
            for chunk in response.iter_content(chunk_size=8192):
                f.write(chunk)
        print(f"  Downloaded (NSE): {file_path}")
        return file_path
    except Exception as e:
        print(f"  Error downloading NSE PDF {pdf_url}: {e}")
        return None


def analyze_single_stock(stock_name, scrip_code, deep_dive=True, nse_symbol=""):
    """
    Performs the full analysis pipeline for a single stock.
    Supports both BSE (scrip_code) and NSE (nse_symbol) announcements.
    """
    print(f"\n{'='*40}\nAnalyzing {stock_name} (BSE: {scrip_code}, NSE: {nse_symbol})...\n{'='*40}")

    announcements = []

    # Prefer NSE over BSE to avoid duplicate data for dual-listed companies
    if nse_symbol:
        nse_announcements = fetch_nse_announcements(nse_symbol)
        announcements.extend(nse_announcements)
        print(f"  NSE: {len(nse_announcements)} announcements (primary source)")
    elif scrip_code:
        # Only use BSE if NSE symbol is not available (BSE-only stocks)
        bse_announcements = fetch_historical_announcements(scrip_code)
        announcements.extend(bse_announcements)
        print(f"  BSE: {len(bse_announcements)} announcements (no NSE symbol, using BSE)")

    collected_texts = []

    if not announcements:
        print(f"No announcements for {stock_name}")
        # Still try to generate report from screener + web data
    else:
        print(f"Found {len(announcements)} total announcements. Filtering for key documents...")

    # 1. Fetch External Intelligence (Web Search)
    print("Fetching External Web Intelligence (News & Sentiment)...")
    web_intel = fetch_web_intel(stock_name)
    collected_texts.append(web_intel)

    # 1b. Fetch Financial Data from Screener.in
    screener_data = fetch_screener_financials(stock_name)
    if screener_data.get("raw_text"):
        collected_texts.append(screener_data["raw_text"])

    relevant_count = 0
    for ann in announcements:
        subject = ann.get("NEWSSUB", "")
        description = ann.get("HEADNAME", "")
        attachment = ann.get("ATTACHMENTNAME", "")
        nse_pdf_url = ann.get("NSE_PDF_URL", "")
        source = ann.get("SOURCE", "BSE")

        # Smart Filter using both subject and description
        if not is_important_document(subject, description):
            continue

        relevant_count += 1
        print(f"  -> [{source}] Found Relevant: {subject[:60]}...")

        if source == "NSE" and nse_pdf_url:
            # Download NSE PDF directly from URL
            fname = nse_pdf_url.split('/')[-1]
            path = download_nse_pdf(nse_pdf_url, fname)
            if path:
                text = extract_text_from_pdf(path)
                if text:
                    collected_texts.append(f"--- Document ({ann.get('NEWS_DT')}): {subject} ---\n{text}")
                try:
                    os.remove(path)
                except:
                    pass
        elif attachment:
            fname = attachment.split('/')[-1]
            path = download_bse_attachment(attachment, fname)
            if path:
                text = extract_text_from_pdf(path)
                if text:
                    collected_texts.append(f"--- Document ({ann.get('NEWS_DT')}): {subject} ---\n{text}")
                try:
                    os.remove(path)
                except:
                    pass

    print(f"Analyzing {len(collected_texts)} significant documents (out of {len(announcements)} total)...")

    # 2. Analyze (Gemini generates everything EXCEPT financial tables)
    print("Generating Deep Dive Report...")
    report = get_collective_summary(stock_name, collected_texts)

    # 3. Inject screener.in financial tables directly into the report
    # Find "## 3. Financial Analysis" and insert tables after it
    screener_section = ""
    if screener_data.get("profit-loss") or screener_data.get("balance-sheet") or screener_data.get("cash-flow"):
        screener_section += "\n\n### 3.1 Consolidated Profit & Loss Statement (Rs. in Crores)\n"
        screener_section += "(Source: Screener.in)\n\n"
        screener_section += screener_data.get("profit-loss", "Data not available") + "\n\n"
        screener_section += "### 3.2 Consolidated Balance Sheet (Rs. in Crores)\n"
        screener_section += "(Source: Screener.in)\n\n"
        screener_section += screener_data.get("balance-sheet", "Data not available") + "\n\n"
        screener_section += "### 3.3 Consolidated Cash Flow Statement (Rs. in Crores)\n"
        screener_section += "(Source: Screener.in)\n\n"
        screener_section += screener_data.get("cash-flow", "Data not available") + "\n"

    if screener_section:
        # Insert right after "## 3. Financial Analysis" and its commentary
        import re as re_module
        # Find Section 3 heading and insert tables before Section 4
        pattern = r'(## 3\. Financial Analysis.*?)(\n## 4\.)'
        match = re_module.search(pattern, report, re_module.DOTALL)
        if match:
            report = report[:match.end(1)] + screener_section + report[match.start(2):]
        else:
            # Fallback: just append before Section 4 or at end
            if "## 4." in report:
                report = report.replace("## 4.", screener_section + "\n## 4.", 1)
            else:
                report += screener_section

    # 4. Save as DOCX
    docx_path = generate_word_report(stock_name, report)

    return docx_path, len(collected_texts)

def main():
    setup_directories()
    stocks = load_stocks()

    if not stocks:
        print("No stocks found in stocks.json.")
        return

    for stock in stocks:
        name = stock['name']
        scrip = stock['scrip_code']
        analyze_single_stock(name, scrip)

if __name__ == "__main__":
    main()
